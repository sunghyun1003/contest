from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

README_PATH = "/home/user/contest/README_통합_마케팅_애널리틱스.txt"
with open(README_PATH, encoding="utf-8") as f:
    content = f.read()

# ─── DOCX ────────────────────────────────────────────────────────────────────
doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xf9, 0x73, 0x16)

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

# 제목
t = doc.add_heading('자동차CM 통합 마케팅 애널리틱스', 0)
t.runs[0].font.color.rgb = RGBColor(0xf9, 0x73, 0x16)
t.runs[0].font.name = '맑은 고딕'
t.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

add_para(doc, '사내 AI 바이브코딩 경진대회 제출작 안내 문서', bold=False)
doc.add_paragraph()

# 배포 URL 강조 박스 효과 (굵게)
url_para = doc.add_paragraph()
url_run = url_para.add_run('배포 URL: https://sunghyun1003.github.io/contest/')
url_run.bold = True
url_run.font.size = Pt(11)
url_run.font.color.rgb = RGBColor(0x16, 0x56, 0xc4)
url_run.font.name = '맑은 고딕'
url_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()

sections_raw = content.split('\n\n')
current_section = []

for line in content.split('\n'):
    stripped = line.strip()
    # 섹션 헤더 감지 (숫자. 으로 시작)
    import re
    if re.match(r'^\d+\. ', stripped) and not stripped.startswith('1)') and not stripped.startswith('2)'):
        if stripped == '자동차CM 통합 마케팅 애널리틱스 README':
            continue
        if stripped.startswith('='):
            continue
        add_heading(doc, stripped, level=2)
    elif stripped.startswith('---'):
        continue
    elif stripped.startswith('==='):
        continue
    elif stripped == '자동차CM 통합 마케팅 애널리틱스 README':
        continue
    elif stripped == '':
        doc.add_paragraph()
    elif stripped.startswith('- ') or stripped.startswith('  - '):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(stripped.lstrip('- ').lstrip())
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    elif re.match(r'^\d+\)', stripped):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(stripped[3:].strip() if len(stripped) > 2 else stripped)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    elif stripped in ['주요 기능', '사용 팁', '주요 사용 방식', '탭별 기능', '목적', '권장 브라우저: Chrome 또는 Microsoft Edge',
                      '접속 방식: URL 직접 입력 또는 링크 클릭', '인터넷 연결: 필수 (GitHub Pages를 통해 호스팅됨)']:
        add_para(doc, stripped, bold=True)
    else:
        add_para(doc, stripped)

docx_path = "/home/user/contest/README_통합_마케팅_애널리틱스.docx"
doc.save(docx_path)
print(f"DOCX 저장 완료: {docx_path}")

# ─── PDF ─────────────────────────────────────────────────────────────────────
# 한글 폰트 찾기
import glob as glob_mod
font_candidates = [
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
]
korean_font = None
for fc in font_candidates:
    if os.path.exists(fc):
        korean_font = fc
        break

if not korean_font:
    found = glob_mod.glob("/usr/share/fonts/**/*.ttf", recursive=True)
    if found:
        korean_font = found[0]

font_name = "Korean"
if korean_font:
    try:
        pdfmetrics.registerFont(TTFont(font_name, korean_font))
    except Exception:
        font_name = "Helvetica"
else:
    font_name = "Helvetica"

pdf_path = "/home/user/contest/README_통합_마케팅_애널리틱스.pdf"
doc_pdf = SimpleDocTemplate(
    pdf_path,
    pagesize=A4,
    leftMargin=20*mm, rightMargin=20*mm,
    topMargin=20*mm, bottomMargin=20*mm
)

styles = getSampleStyleSheet()
title_style = ParagraphStyle('KTitle', fontName=font_name, fontSize=18, textColor=colors.HexColor('#f97316'),
                              spaceAfter=6, leading=26)
h2_style = ParagraphStyle('KH2', fontName=font_name, fontSize=13, textColor=colors.HexColor('#f97316'),
                           spaceAfter=4, spaceBefore=12, leading=20)
url_style = ParagraphStyle('KURL', fontName=font_name, fontSize=11, textColor=colors.HexColor('#1656c4'),
                            spaceAfter=8, leading=16)
body_style = ParagraphStyle('KBody', fontName=font_name, fontSize=9.5, leading=16, spaceAfter=2)
bullet_style = ParagraphStyle('KBullet', fontName=font_name, fontSize=9.5, leading=16,
                               leftIndent=12, bulletIndent=0, spaceAfter=2)
bold_style = ParagraphStyle('KBold', fontName=font_name, fontSize=9.5, leading=16, spaceAfter=2)

story = []
story.append(Paragraph('자동차CM 통합 마케팅 애널리틱스', title_style))
story.append(Paragraph('사내 AI 바이브코딩 경진대회 제출작 안내 문서', body_style))
story.append(Spacer(1, 6*mm))
story.append(Paragraph('<b>배포 URL: https://sunghyun1003.github.io/contest/</b>', url_style))
story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
story.append(Spacer(1, 4*mm))

bold_keywords = ['주요 기능', '사용 팁', '주요 사용 방식', '탭별 기능', '목적',
                 '권장 브라우저', '접속 방식', '인터넷 연결', '[참고]']

import re
for line in content.split('\n'):
    stripped = line.strip()
    if stripped == '자동차CM 통합 마케팅 애널리틱스 README':
        continue
    if re.match(r'^=+$', stripped) or re.match(r'^-+$', stripped):
        continue
    if stripped == '':
        story.append(Spacer(1, 3*mm))
        continue
    if re.match(r'^\d+\. ', stripped) and not re.match(r'^\d+\)', stripped):
        story.append(Paragraph(stripped, h2_style))
    elif stripped.startswith('- ') or stripped.startswith('  - '):
        text = stripped.lstrip('- ').lstrip()
        story.append(Paragraph(f'• {text}', bullet_style))
    elif re.match(r'^\d+\)', stripped):
        story.append(Paragraph(stripped, bullet_style))
    elif any(stripped.startswith(k) for k in bold_keywords):
        story.append(Paragraph(f'<b>{stripped}</b>', bold_style))
    else:
        story.append(Paragraph(stripped, body_style))

doc_pdf.build(story)
print(f"PDF 저장 완료: {pdf_path}")
